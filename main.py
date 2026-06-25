#!/usr/bin/env python3
"""
巴西资讯日报 - 多源采集 + QQ邮箱推送（HTML正文 + Word附件）+ 微信公众号自动发布
GitHub Actions 定时执行（每天UTC 1:00 = 北京时间9:00）
"""
import os, sys, json, ssl, smtplib, tempfile, hashlib, re
from datetime import datetime, timezone, timedelta, date
from email.message import EmailMessage
from urllib.request import Request, urlopen
from urllib.parse import quote
import html as _html

# ---- python-docx for Word attachment ----
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ===== 配置 =====
SMTP_HOST = "smtp.qq.com"
SMTP_PORT = 465
SENDER = "847988716@qq.com"
RECEIVER = "847988716@qq.com"
AUTH_CODE = os.environ.get("QQ_MAIL_AUTH_CODE", "")

# WeChat Official Account credentials (optional - only set if you want auto-publish)
WECHAT_APPID = os.environ.get("WECHAT_APPID", "")
WECHAT_APPSECRET = os.environ.get("WECHAT_APPSECRET", "")
WECHAT_ENABLED = bool(WECHAT_APPID and WECHAT_APPSECRET)

# Beijing timezone
BJ_TZ = timezone(timedelta(hours=8))

# 今日日期过滤字符串（格式 "YYYY-MM-DD"，在 main() 中赋值）
# 用于 Google News after: 操作符，确保只抓当天新鲜新闻
_TODAY_FILTER: str = ""

# ===== 中文翻译模块 =====

def _try_translate_google(text):
    """Google Translate（GitHub Actions可用，国内会超时）"""
    try:
        from deep_translator import GoogleTranslator
        return GoogleTranslator(source='auto', target='zh-CN').translate(text[:800])
    except Exception:
        return None

def _try_translate_mymemory(text):
    """MyMemory 免费翻译（无需API Key，国内可能可用）"""
    try:
        from deep_translator import MyMemoryTranslator
        return MyMemoryTranslator(source='auto', target='zh-CN').translate(text[:800])
    except Exception:
        return None

def translate_to_chinese(text):
    """将非中文文本翻译为简体中文，三引擎后备"""
    if not text or len(text.strip()) < 3:
        return text
    # 跳过已是中文的内容（含中文比例超过30%视为中文）
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    if chinese_chars / max(len(text), 1) > 0.3:
        return text
    
    # 依次尝试：Google → MyMemory
    for engine, name in [(_try_translate_google, "Google"), (_try_translate_mymemory, "MyMemory")]:
        try:
            result = engine(text)
            if result and result != text and len(result.strip()) > 3:
                return result
        except Exception:
            continue
    
    # 全部失败 → 保留原文（外部源的文章将在GitHub Actions上正常翻译）
    return text

def translate_article(article):
    """翻译单条资讯的标题和摘要为中文"""
    orig_title = article.get("title", "")
    orig_summary = article.get("summary", "")
    article["title_zh"] = translate_to_chinese(orig_title)
    article["summary_zh"] = translate_to_chinese(orig_summary)
    return article

def is_chinese(text):
    """判断文本是否主要为中文"""
    if not text:
        return False
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    return chinese_chars / max(len(text), 1) > 0.3

# ===== 日期判断工具 =====
def parse_rss_date(date_str):
    """Parse RSS pubDate / standard date strings, return datetime in BJT"""
    if not date_str:
        return None
    text = date_str.strip()  # 不截断，保留完整日期字符串以正确解析
    # RSS standard: "Mon, 22 Jun 2026 14:30:00 GMT"
    for fmt in [
        "%a, %d %b %Y %H:%M:%S %Z",
        "%a, %d %b %Y %H:%M:%S",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d/%m/%Y %H:%M",
    ]:
        try:
            dt = datetime.strptime(text, fmt)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(BJ_TZ)
        except (ValueError, IndexError):
            continue
    return None

def is_today_article(article, bj_today):
    """判断文章是否为近期发布的真实新闻（最近2天内）"""
    t = article.get("title", "").lower()
    # 排除明显是网站简介/指引类的内容（不含具体新闻事实）
    generic_kw = ["请访问", "请查看", "定期发布", "最新一期",
                  "欢迎关注", "覆盖", "提供",
                  "latest", "visit", "check out"]
    has_generic = any(kw in t for kw in generic_kw)
    
    # 尝试从时间字段解析日期
    time_str = article.get("time", "")
    parsed = parse_rss_date(time_str)
    
    if parsed:
        # 接受最近2天内的文章（Google News有时会返回昨天的）
        days_diff = (bj_today - parsed.date()).days
        is_recent = 0 <= days_diff <= 2
        if is_recent:
            return True
        elif has_generic:
            return False
        else:
            return False
    else:
        # 无时间字段 → 如果有兜底特征则排除，否则保守保留（来自Google News等实时源）
        if has_generic:
            return False
        s = article.get("summary", "").lower()
        if any(kw in s for kw in generic_kw):
            return False
        # 无法判断日期，保留（Google News文章通常有时效性）
        return True

# ===== HTML富文本生成器 =====
def build_html(articles, bj_now):
    """生成巴西资讯日报HTML邮件（全中文）"""
    date_str = bj_now.strftime("%Y年%m月%d日")
    date_full = bj_now.strftime("%Y-%m-%d %H:%M")
    
    src_count = len(set(a["source"] for a in articles))
    total = len(articles)
    
    # Group by category
    cats = {}
    for a in articles:
        cat = a.get("category", "综合")
        cats.setdefault(cat, []).append(a)
    
    items_html = ""
    for cat, items in cats.items():
        items_html += f'<tr><td colspan="2" style="background:#009739;color:#fff;padding:10px 16px;font-size:15px;font-weight:bold;text-align:center;letter-spacing:2px;">📌 {cat}</td></tr>'
        for i, item in enumerate(items, 1):
            url = item.get("url", "#")
            # 优先使用中文翻译
            display_title = item.get("title_zh") or item["title"]
            display_summary = item.get("summary_zh") or item["summary"]
            items_html += f"""<tr style="background:{('#fff' if i%2==1 else '#f8fff5')};">
<td style="padding:14px 16px;vertical-align:top;width:28px;color:#009739;font-weight:bold;font-size:16px;">{i}</td>
<td style="padding:14px 16px;">
<div style="font-size:16px;font-weight:bold;color:#1a3a1a;margin-bottom:6px;line-height:1.5;">{display_title}</div>
<div style="font-size:12px;color:#888;margin-bottom:8px;">📡 {item['source']} &nbsp;|&nbsp; 🕐 {item.get('time','')}</div>
<div style="font-size:14px;color:#333;line-height:1.8;margin-bottom:10px;">{display_summary}</div>
<a href="{url}" target="_blank" style="display:inline-block;padding:6px 18px;background:#009739;color:#fff;text-decoration:none;border-radius:4px;font-size:13px;font-weight:bold;">📖 阅读原文 →</a>
</td></tr>"""
    
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"></head>
<body style="font-family:'Microsoft YaHei',sans-serif;margin:0;padding:0;background:#f0f4f0;">

<table width="100%" cellpadding="0" cellspacing="0" style="background:#f0f4f0;padding:20px 0;">
<tr><td align="center">
<table width="680" cellpadding="0" cellspacing="0" style="background:#fff;border-radius:8px;overflow:hidden;">

<tr><td style="background:linear-gradient(135deg,#009739,#00b44a);padding:28px 30px;text-align:center;">
<div style="font-size:12px;color:#a3e6b3;letter-spacing:2px;">巴西每日财经简报</div>
<div style="font-size:26px;font-weight:bold;color:#fff;margin:6px 0;">巴西资讯日报</div>
<div style="font-size:14px;color:#d4f5db;">{date_str}</div>
</td></tr>

<tr><td style="padding:20px 24px;background:#eaf7ed;font-size:13px;color:#336;">
  当日共采集 {total} 条资讯，覆盖 {src_count} 个信息源 &nbsp;|&nbsp; 生成时间 {date_full}（北京时间）<br>
  🌐 <a href="https://lucane3.github.io/brazil-news-bot/" style="color:#009739;font-weight:bold;">国内直连查看完整日报 →</a>
</td></tr>

<tr><td style="padding:4px 0;">
<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;">
{items_html}
</table></td></tr>

<tr><td style="padding:16px 24px;text-align:center;font-size:11px;color:#aaa;border-top:1px solid #e8e8e8;">
巴西资讯日报 · 自动生成
</td></tr>

</table>
</td></tr></table>
</body></html>"""

# ===== 公众号粘贴专用HTML生成器 =====

def build_wechat_publish_html(articles, bj_now):
    """生成可直接复制粘贴到微信公众号编辑器的HTML文件（与邮件格式完全一致）"""
    date_str = bj_now.strftime("%Y年%m月%d日")
    date_full = bj_now.strftime("%Y-%m-%d %H:%M")
    
    src_count = len(set(a["source"] for a in articles))
    total = len(articles)
    
    cats = {}
    for a in articles:
        cat = a.get("category", "综合")
        cats.setdefault(cat, []).append(a)
    
    items_html = ""
    for cat, items in cats.items():
        items_html += f'              <tr>\n                <td colspan="2" style="background: #009739; color: #ffffff; padding: 10px 16px; font-size: 15px; font-weight: bold; text-align: center; letter-spacing: 2px;">📌 {cat}</td>\n              </tr>\n'
        for i, item in enumerate(items, 1):
            url = item.get("url", "#")
            display_title = item.get("title_zh") or item["title"]
            display_summary = item.get("summary_zh") or item.get("summary", "")
            bg_color = "#ffffff" if i % 2 == 1 else "#f8fff5"
            items_html += f'''              <tr style="background: {bg_color};">
                <td style="padding: 14px 16px; vertical-align: top; width: 28px; color: #009739; font-weight: bold; font-size: 16px;">{i}</td>
                <td style="padding: 14px 16px;">
                  <div style="font-size: 16px; font-weight: bold; color: #1a3a1a; margin-bottom: 6px; line-height: 1.5;">{display_title}</div>
                  <div style="font-size: 12px; color: #888888; margin-bottom: 8px;">📡 {item['source']} &nbsp;|&nbsp; 🕐 {item.get('time','')}</div>
                  <div style="font-size: 14px; color: #333333; line-height: 1.8; margin-bottom: 10px;">{display_summary}</div>
                  <a href="{url}" target="_blank" style="display: inline-block; padding: 6px 18px; background: #009739; color: #ffffff; text-decoration: none; font-size: 13px; font-weight: bold;">📖 阅读原文 →</a>
                </td>
              </tr>
'''
    
    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>巴西资讯日报 - 公众号粘贴专用 | {date_str}</title>
<style>
  body {{ font-family: 'Microsoft YaHei', sans-serif; margin: 0; padding: 20px; background: #e8e8e8; }}
  .preview-wrap {{ max-width: 680px; margin: 0 auto; }}
</style>
</head>
<body>
<div class="preview-wrap">
<table width="100%" cellpadding="0" cellspacing="0" style="background: #f0f4f0; padding: 20px 0;">
  <tr>
    <td align="center">
      <table width="680" cellpadding="0" cellspacing="0" style="background: #ffffff;">
        <tr>
          <td style="background: #009739; padding: 28px 30px; text-align: center;">
            <div style="font-size: 12px; color: #a3e6b3; letter-spacing: 2px;">每日巴西财经简报</div>
            <div style="font-size: 26px; font-weight: bold; color: #ffffff; margin: 6px 0;">巴西资讯日报</div>
            <div style="font-size: 14px; color: #d4f5db;">{date_str}</div>
          </td>
        </tr>
        <tr>
          <td style="padding: 20px 24px; background: #eaf7ed; font-size: 13px; color: #333366;">
            当日共采集 <b>{total}</b> 条资讯，覆盖 <b>{src_count}</b> 个信息源 &nbsp;|&nbsp; 生成时间 {date_full}（北京时间）<br>
            🌐 <a href="https://lucane3.github.io/brazil-news-bot/" target="_blank" style="color: #009739; font-weight: bold;">国内直连查看完整日报 →</a>
          </td>
        </tr>
        <tr>
          <td style="padding: 4px 0;">
            <table width="100%" cellpadding="0" cellspacing="0" style="border-collapse: collapse;">
{items_html}
            </table>
          </td>
        </tr>
        <tr>
          <td style="padding: 16px 24px; text-align: center; font-size: 11px; color: #aaaaaa; border-top: 1px solid #e8e8e8;">
            巴西资讯日报 · 自动生成 | 信息源：Google News / BCB巴西央行 / IBGE统计局 / 新华网 / 南美侨报网 / 巴西联邦税务局
          </td>
        </tr>
      </table>
    </td>
  </tr>
</table>
</div>
</body>
</html>'''


# ===== GitHub Pages 国内可访问页面 =====

def build_github_page(articles, bj_now):
    """生成完整的HTML页面用于GitHub Pages（国内可直连访问）"""
    date_str = bj_now.strftime("%Y年%m月%d日")
    date_full = bj_now.strftime("%Y-%m-%d %H:%M")
    src_count = len(set(a["source"] for a in articles))
    total = len(articles)
    
    # Group by category
    cats = {}
    for a in articles:
        cat = a.get("category", "综合")
        cats.setdefault(cat, []).append(a)
    
    items_html = ""
    for cat, items in cats.items():
        items_html += f'<div class="category"><h2 class="cat-title">📌 {cat}</h2>'
        for i, item in enumerate(items, 1):
            display_title = item.get("title_zh") or item["title"]
            display_summary = item.get("summary_zh") or item.get("summary", "")
            url = item.get("url", "#")
            source = item.get("source", "")
            time_str = item.get("time", "")
            
            items_html += f'''
            <div class="article">
                <div class="article-number">{i}</div>
                <div class="article-content">
                    <h3>{display_title}</h3>
                    <div class="meta">📡 {source} &nbsp;|&nbsp; 🕐 {time_str}</div>
                    <p class="summary">{display_summary}</p>
                    <a href="{url}" target="_blank" class="btn">📖 查看原文 →</a>
                </div>
            </div>'''
            if i < len(items):
                items_html += '<hr class="divider">'
        items_html += '</div>'
    
    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>巴西资讯日报 | {date_str}</title>
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{ font-family:'Microsoft YaHei','PingFang SC',sans-serif; background:#f0f4f0; color:#333; line-height:1.8; }}
.header {{ background:linear-gradient(135deg,#009739,#00b44a); color:#fff; text-align:center; padding:40px 20px; }}
.header h1 {{ font-size:32px; letter-spacing:3px; }}
.header .sub {{ font-size:14px; color:#d4f5db; margin-top:6px; }}
.header .date {{ font-size:18px; color:#fff; margin-top:10px; }}
.stats {{ background:#eaf7ed; text-align:center; padding:14px; font-size:14px; color:#336; }}
.container {{ max-width:800px; margin:0 auto; padding:20px; }}
.category {{ margin-bottom:30px; }}
.cat-title {{ background:#009739; color:#fff; padding:10px 20px; font-size:16px; border-radius:4px; display:inline-block; margin-bottom:16px; }}
.article {{ display:flex; gap:16px; padding:16px; background:#fff; border-radius:6px; margin-bottom:12px; box-shadow:0 1px 4px rgba(0,0,0,0.06); }}
.article-number {{ flex-shrink:0; width:32px; height:32px; background:#009739; color:#fff; border-radius:50%; display:flex; align-items:center; justify-content:center; font-weight:bold; font-size:14px; }}
.article-content {{ flex:1; }}
.article-content h3 {{ font-size:17px; color:#1a3a1a; margin-bottom:4px; }}
.meta {{ font-size:12px; color:#888; margin-bottom:8px; }}
.summary {{ font-size:14px; color:#444; margin-bottom:10px; text-indent:2em; }}
.btn {{ display:inline-block; padding:6px 18px; background:#009739; color:#fff; text-decoration:none; border-radius:4px; font-size:13px; font-weight:bold; }}
.btn:hover {{ background:#007a2e; }}
.divider {{ border:none; border-top:1px dashed #ddd; margin:4px 0 12px 48px; }}
.footer {{ text-align:center; padding:30px; font-size:12px; color:#aaa; border-top:1px solid #e8e8e8; margin-top:20px; }}
.footer a {{ color:#009739; }}
@media (max-width:600px) {{
    .container {{ padding:10px; }}
    .header h1 {{ font-size:24px; }}
    .article {{ flex-direction:column; gap:8px; }}
}}
</style>
</head>
<body>
<div class="header">
    <div class="sub">每日巴西财经简报</div>
    <h1>巴西资讯日报</h1>
    <div class="date">{date_str}</div>
</div>
<div class="stats">
    当日共采集 {total} 条资讯，覆盖 {src_count} 个信息源 &nbsp;|&nbsp; 生成时间 {date_full}（北京时间）
</div>
<div class="container">
{items_html}
</div>
<div class="footer">
    <p>巴西资讯日报 · 自动生成 | 信息源：Google News / BCB巴西央行 / IBGE统计局 / 新华网 / 南美侨报网 / 巴西联邦税务局</p>
    <p>数据来源链接均来自海外媒体，部分可能需要科学上网访问</p>
</div>
</body>
</html>'''

# ===== Word文档生成器（适合微信公众号发布） =====

# 巴西绿主题色
BR_GREEN = RGBColor(0, 151, 57)
BR_YELLOW = RGBColor(253, 217, 0)
BR_DARK = RGBColor(26, 58, 26)
GRAY_TEXT = RGBColor(120, 120, 120)
WHITE = RGBColor(255, 255, 255)
LIGHT_GREEN = RGBColor(240, 250, 240)

def _set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)

def _set_run_font(run, name='Microsoft YaHei', size=11, bold=False, color=None):
    """Set run font properties including East-Asian font"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East-Asian font for CJK
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): name})
    rPr.insert(0, rFonts)

def _add_paragraph_styled(doc, text, style='Normal', bold=False, size=11, color=None, alignment=None, space_after=6):
    """Add a styled paragraph"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    return p

def build_docx(articles, bj_now):
    """生成全中文Word文档（适合微信公众号发布）"""
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    date_str = bj_now.strftime("%Y年%m月%d日")
    date_short = bj_now.strftime("%Y-%m-%d")
    
    # ===== 封面标题区 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    run = p.add_run("巴西资讯日报")
    _set_run_font(run, size=36, bold=True, color=BR_GREEN)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    run = p.add_run("每日巴西财经简报")
    _set_run_font(run, size=14, bold=False, color=GRAY_TEXT)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str)
    _set_run_font(run, size=16, bold=True, color=BR_DARK)
    
    # Green divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(20)
    run = p.add_run("━" * 36)
    _set_run_font(run, size=8, bold=False, color=BR_GREEN)
    
    # 摘要统计
    src_count = len(set(a["source"] for a in articles))
    total = len(articles)
    summary = f"收录 {total} 条资讯  ·  覆盖 {src_count} 个信息源  ·  {date_short}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(16)
    run = p.add_run(summary)
    _set_run_font(run, size=10, bold=False, color=GRAY_TEXT)
    
    # ===== 分类排版：每条资讯 =====
    cats = {}
    for a in articles:
        cat = a.get("category", "综合")
        cats.setdefault(cat, []).append(a)
    
    for cat, items in cats.items():
        # 分类标题 (绿色背景条)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(14)
        pf.space_after = Pt(8)
        run = p.add_run(f"  📌 {cat}  ")
        _set_run_font(run, size=13, bold=True, color=WHITE)
        # Add shading via XML
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:fill'): '009739',
            qn('w:val'): 'clear',
        })
        pPr.append(shd)
        
        for i, item in enumerate(items, 1):
            # 标题（优先中文译文）
            display_title = item.get("title_zh") or item["title"]
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(2)
            run = p.add_run(f"{i}. {display_title}")
            _set_run_font(run, size=14, bold=True, color=BR_DARK)
            
            # 来源 + 时间
            src_line = f"  来源：{item['source']}"
            if item.get('time'):
                src_line += f"  |  {item['time']}"
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(4)
            run = p.add_run(src_line)
            _set_run_font(run, size=9, bold=False, color=GRAY_TEXT)
            
            # 正文摘要（优先中文译文）
            summary_text = item.get("summary_zh") or item.get("summary", "")
            if summary_text:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_after = Pt(6)
                pf.first_line_indent = Cm(0.7)
                run = p.add_run(summary_text)
                _set_run_font(run, size=11, bold=False, color=None)
            
            # 原文链接
            url = item.get("url", "")
            if url:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_after = Pt(10)
                run = p.add_run(f"  阅读原文：{url}")
                _set_run_font(run, size=9, bold=False, color=RGBColor(0, 151, 57))
            
            # 分隔线 (除了最后一条)
            if i < len(items):
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_after = Pt(2)
                run = p.add_run("—" * 40)
                _set_run_font(run, size=6, bold=False, color=RGBColor(200, 200, 200))
    
    # ===== 页脚 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(4)
    run = p.add_run("━" * 36)
    _set_run_font(run, size=8, bold=False, color=BR_GREEN)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("巴西资讯日报 · 自动生成")
    _set_run_font(run, size=9, bold=False, color=GRAY_TEXT)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("信息源：Google News(多语种) / 巴西央行BCB / IBGE统计局 / 新华网葡语版 / 南美侨报网 / 巴西联邦税务局")
    _set_run_font(run, size=8, bold=False, color=RGBColor(170, 170, 170))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("适合微信公众号平台发布 · 可直接复制排版")
    _set_run_font(run, size=8, bold=False, color=RGBColor(200, 200, 200))
    
    return doc


# ===== 微信公众号自动发布模块 =====
def wechat_get_token(appid, appsecret):
    """获取微信公众号 access_token（有效期2小时）"""
    url = f"https://api.weixin.qq.com/cgi-bin/token?grant_type=client_credential&appid={appid}&secret={appsecret}"
    data = _json_fetch(url, timeout=10)
    if data and "access_token" in data:
        return data["access_token"]
    print(f"  [WECHAT WARN] Failed to get access_token: {data}")
    return None

def wechat_upload_cover_image(access_token, docx_path):
    """将Word文档首页截图思路 → 改为生成纯色封面图并上传为永久素材"""
    # WeChat requires a cover image (thumb_media_id) for each article
    # Generate a simple green cover PNG using PIL if available, else use a placeholder approach
    try:
        from PIL import Image, ImageDraw, ImageFont
        import io
        img = Image.new("RGB", (900, 500), "#009739")
        draw = ImageDraw.Draw(img)
        # Draw title text
        draw.rectangle([20, 20, 880, 480], outline="#ffffff", width=2)
        try:
            font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 48)
        except Exception:
            font = ImageFont.load_default()
        draw.text((450, 200), "巴西资讯日报", fill="white", font=font, anchor="mm")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        img_bytes = buf.read()
    except ImportError:
        # Fallback: simple 1x1 green pixel PNG (minimal valid PNG)
        img_bytes = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x01\x90\x00\x00\x01\x90\x08\x02\x00\x00\x00h6\xc2\x99\x00\x00\x00\x01sRGB\x00\xae\xce\x1c\xe9\x00\x00\x00&IDATx\xda\xed\xc1\x01\x01\x00\x00\x00\x82 \xff\xafnH@\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\xff\xff\x03\x00\xf5%\x0f\x0e\x006\xbe&\x96\x00\x00\x00\x00IEND\xaeB`\x82'
    
    # Upload as permanent material
    boundary = "----WebKitFormBoundary7MA4YWxkTrZu0gW"
    body = (
        f'--{boundary}\r\n'
        f'Content-Disposition: form-data; name="media"; filename="cover.png"\r\n'
        f'Content-Type: image/png\r\n\r\n'
    ).encode("utf-8") + img_bytes + f'\r\n--{boundary}--\r\n'.encode("utf-8")
    
    url = f"https://api.weixin.qq.com/cgi-bin/material/add_material?access_token={access_token}&type=image"
    req = Request(url, data=body)
    req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
    req.add_header("User-Agent", "BrazilNewsBot/1.0")
    try:
        with urlopen(req, timeout=20) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            if "media_id" in result:
                return result["media_id"]
            print(f"  [WECHAT WARN] Upload cover failed: {result}")
    except Exception as e:
        print(f"  [WECHAT WARN] Upload cover error: {e}")
    return None

def wechat_build_article_html(deduped):
    """构建适用于微信公众号的富文本HTML（全中文，兼容微信客户端）"""
    parts = ['<section style="padding: 10px 16px; max-width: 100%%; box-sizing: border-box;">']
    
    # Title
    parts.append('<p style="text-align:center;"><span style="font-size:22px;color:#009739;font-weight:bold;">巴西资讯日报</span></p>')
    parts.append('<p style="text-align:center;"><span style="font-size:13px;color:#888888;">每日巴西财经简报</span></p>')
    
    # Divider
    parts.append('<p style="text-align:center;"><span style="color:#009739;">━━━━━━━━━━</span></p>')
    
    cats = {}
    for a in deduped:
        cats.setdefault(a.get("category", "综合"), []).append(a)
    
    for cat, items in cats.items():
        parts.append(f'<p style="background-color:#009739;color:#ffffff;padding:6px 12px;font-size:15px;font-weight:bold;border-radius:3px;">📌 {cat}</p>')
        for i, item in enumerate(items, 1):
            # 优先使用中文译文
            title = (item.get("title_zh") or item["title"]).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            summary = (item.get("summary_zh") or item.get("summary", "")).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # 限制摘要长度
            summary = summary[:300] if len(summary) > 300 else summary
            source = item.get("source", "")
            url = item.get("url", "")
            
            parts.append(f'<p style="margin-top:12px;"><strong style="font-size:16px;color:#1a3a1a;">{i}. {title}</strong></p>')
            parts.append(f'<p style="font-size:11px;color:#888888;">来源：{source}</p>')
            if summary:
                parts.append(f'<p style="font-size:14px;color:#333333;line-height:1.8;">{summary}</p>')
            if url:
                parts.append(f'<p><a href="{url}" style="color:#009739;font-size:13px;">📖 阅读原文 →</a></p>')
            if i < len(items):
                parts.append('<p style="border-bottom:1px dashed #e0e0e0;"></p>')
    
    parts.append(f'<p style="text-align:center;margin-top:20px;font-size:10px;color:#aaaaaa;">巴西资讯日报 · 自动生成 | 共{len(deduped)}条资讯</p>')
    parts.append('</section>')
    return "\n".join(parts)

def wechat_publish(deduped, bj_now):
    """发布到微信公众号草稿箱（需设置 WECHAT_APPID / WECHAT_APPSECRET）"""
    if not WECHAT_ENABLED:
        print("[WECHAT] Not configured, skipping publish.")
        return None
    
    print("[WECHAT] Getting access token...")
    token = wechat_get_token(WECHAT_APPID, WECHAT_APPSECRET)
    if not token:
        return None
    
    print("[WECHAT] Uploading cover image...")
    cover_media_id = wechat_upload_cover_image(token, None)
    if not cover_media_id:
        print("[WECHAT WARN] No cover image, publish aborted.")
        return None
    
    print("[WECHAT] Building article HTML...")
    content_html = wechat_build_article_html(deduped)
    digest = f"每日巴西资讯汇总，收录{len(deduped)}条，覆盖巴西央行/IBGE/Reuters/新华网/南美侨报/税务局等{len(set(a['source'] for a in deduped))}个源"
    if len(digest) > 120:
        digest = digest[:117] + "..."
    
    date_str = bj_now.strftime("%Y-%m-%d")
    draft_payload = {
        "articles": [{
            "title": f"巴西资讯日报 | {date_str}",
            "author": "巴西资讯日报",
            "digest": digest,
            "content": content_html,
            "content_source_url": "",
            "thumb_media_id": cover_media_id,
            "need_open_comment": 0,
            "only_fans_can_comment": 0,
        }]
    }
    
    print("[WECHAT] Creating draft...")
    url = f"https://api.weixin.qq.com/cgi-bin/draft/add?access_token={token}"
    req = Request(url, data=json.dumps(draft_payload, ensure_ascii=False).encode("utf-8"))
    req.add_header("Content-Type", "application/json; charset=utf-8")
    req.add_header("User-Agent", "BrazilNewsBot/1.0")
    
    try:
        with urlopen(req, timeout=15) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            if "media_id" in result:
                draft_media_id = result["media_id"]
                print(f"[WECHAT] Draft created: {draft_media_id}")
                
                # Publish the draft
                pub_url = f"https://api.weixin.qq.com/cgi-bin/freepublish/submit?access_token={token}"
                pub_payload = json.dumps({"media_id": draft_media_id}, ensure_ascii=False).encode("utf-8")
                pub_req = Request(pub_url, data=pub_payload)
                pub_req.add_header("Content-Type", "application/json; charset=utf-8")
                pub_req.add_header("User-Agent", "BrazilNewsBot/1.0")
                with urlopen(pub_req, timeout=15) as pub_resp:
                    pub_result = json.loads(pub_resp.read().decode("utf-8"))
                    print(f"[WECHAT] Publish result: {pub_result}")
                    if "publish_id" in pub_result:
                        print(f"[WECHAT] ✅ Published successfully! publish_id={pub_result['publish_id']}")
                        return pub_result.get("publish_id")
                    else:
                        print(f"[WECHAT WARN] Publish may have failed: {pub_result}")
            else:
                print(f"[WECHAT WARN] Draft creation failed: {result}")
    except Exception as e:
        print(f"[WECHAT] Error: {e}")
    return None


# ===== 多源采集模块 =====

def _fetch(url, headers=None, timeout=15):
    """HTTP GET with browser-like headers for better success rate"""
    req = Request(url, headers=headers or {})
    req.add_header("User-Agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")
    req.add_header("Accept", "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8")
    req.add_header("Accept-Language", "pt-BR,pt;q=0.9,en;q=0.8,zh-CN;q=0.7,zh;q=0.6")
    req.add_header("Cache-Control", "no-cache")
    try:
        with urlopen(req, timeout=timeout) as resp:
            raw = resp.read()
            for enc in ["utf-8", "latin-1", "iso-8859-1"]:
                try:
                    return raw.decode(enc)
                except (UnicodeDecodeError, LookupError):
                    continue
            return raw.decode("utf-8", errors="replace")
    except Exception as e:
        print(f"  [WARN] Fetch failed: {url[:100]} -> {e}")
        return ""

def _json_fetch(url, timeout=15):
    """Fetch and parse JSON response"""
    text = _fetch(url, timeout=timeout)
    if text:
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
    return None

def _extract(text, start, end):
    """Extract substring between start and end markers"""
    try:
        i = text.index(start) + len(start)
        j = text.index(end, i)
        return text[i:j]
    except (ValueError, IndexError):
        return ""

def _clean(text):
    """Strip HTML tags, CDATA, entities, and normalize whitespace"""
    # Remove CDATA wrappers: <![CDATA[...]]>
    text = re.sub(r'<!\[CDATA\[(.*?)\]\]>', r'\1', text, flags=re.DOTALL)
    # Remove HTML comments
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    # Decode HTML entities
    text = _html.unescape(text)
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def _parse_rss_items(xml_text, max_items=5):
    """Parse RSS XML items into article dicts. Returns empty list if no valid items."""
    if not xml_text or "<item>" not in xml_text:
        return []
    
    items = xml_text.split("<item>")[1:]
    results = []
    for item in items[:max_items]:
        title = _extract(item, "<title>", "</title>")
        link = _extract(item, "<link>", "</link>")
        pubdate = _extract(item, "<pubDate>", "</pubDate>")
        desc = _extract(item, "<description>", "</description>")
        # Extract source name from <source> tag
        src_block = _extract(item, "<source", "</source>")
        src_name = _extract(src_block + ">", "<", "") if src_block else ""
        
        if not title or not link:
            continue
        
        clean_title = _clean(title)
        clean_desc = _clean(desc)[:500] if desc else clean_title
        
        # Google News format: "Actual Title - Source Name"
        display_source = src_name or "Google News"
        if not src_name and " - " in clean_title:
            parts = clean_title.rsplit(" - ", 1)
            if len(parts) == 2 and len(parts[1]) < 50:
                clean_title = parts[0]
                display_source = parts[1]
        
        results.append({
            "title": clean_title,
            "summary": clean_desc,
            "source": display_source,
            "category": "",
            "time": (pubdate or "")[:25].strip(),
            "url": link,
        })
    return results

def _google_news(query, hl="pt-BR", gl="BR", max_items=5):
    """Fetch real news articles from Google News RSS for a given search query.
    Returns empty list if no results (no static fallback descriptions)."""
    ceid = f"{gl}:{hl.split('-')[0]}"
    q = quote(query)
    
    urls = [
        f"https://news.google.com/rss/search?q={q}&hl={hl}&gl={gl}&ceid={ceid}",
        f"https://news.google.com/rss/search?q={q}&hl={hl}&gl={gl}",
    ]
    
    for url in urls:
        text = _fetch(url, timeout=20)
        items = _parse_rss_items(text, max_items=max_items)
        if items:
            return items
    return []

def _bing_news(query, hl="zh-Hans", max_items=5):
    """Fetch news from Bing News RSS (国内可访问的 Google News 替代方案).
    Returns empty list if no results."""
    # Map language codes to Bing's setlang parameter
    lang_map = {
        "pt-BR": "pt", "pt": "pt",
        "en": "en", "en-US": "en",
        "zh-CN": "zh-Hans", "zh": "zh-Hans",
    }
    setlang = lang_map.get(hl, hl.split("-")[0] if "-" in hl else hl)
    
    q = quote(query)
    # Bing News RSS endpoint (works from China mainland)
    urls = [
        f"https://www.bing.com/news/search?q={q}&format=RSS&setlang={setlang}&first=1",
        f"https://www.bing.com/news/search?q={q}&format=RSS",
    ]
    
    for url in urls:
        text = _fetch(url, timeout=20, 
                      headers={"Accept": "application/rss+xml, application/xml, text/xml,*/*"})
        if not text or "<item>" not in text:
            continue
        # Bing RSS uses slightly different tags sometimes
        items_xml = text.split("<item>")[1:]
        results = []
        for item in items_xml[:max_items]:
            title = _extract(item, "<title>", "</title>")
            link = _extract(item, "<link>", "</link>")
            pubdate = _extract(item, "<pubDate>", "</pubDate>") or _extract(item, "<dc:date>", "</dc:date>")
            desc = _extract(item, "<description>", "</description>")
            # Bing wraps source in <source> tag with text content
            src = _extract(item, "<source", "</source>")
            src_name = _extract(src + ">", "<", "/>") if src else "Bing News"
            
            if not title or not link:
                continue
            clean_title = _clean(title)
            clean_desc = _clean(desc)[:500] if desc else clean_title
            
            # Bing titles sometimes have "Title - Source"
            display_source = _clean(src_name) if src_name else "Bing News"
            if " - " in clean_title and len(clean_title.rsplit(" - ", 1)[-1]) < 40:
                parts = clean_title.rsplit(" - ", 1)
                clean_title = parts[0]
                if parts[1].strip() != display_source:
                    display_source = parts[1].strip()
            
            results.append({
                "title": clean_title,
                "summary": clean_desc,
                "source": display_source,
                "category": "",
                "time": (pubdate or "")[:25].strip(),
                "url": link,
            })
        if results:
            return results
    return []

def _news_chain(primary_func, gn_queries, bing_queries, category, max_items=3):
    """通用新闻采集链路：主站 → Google News（今日优先）→ Google News（无日期约束）→ Bing News"""
    # 1) Try primary source
    try:
        results = primary_func()
    except Exception as e:
        print(f"  [WARN] Primary source failed: {e}")
        results = []
    if results:
        for r in results:
            r.setdefault("category", category)
        return results
    
    # 2) Google News with today-only filter（优先获取当天最新内容）
    if _TODAY_FILTER:
        for q, hl, gl in gn_queries:
            q_today = f"{q} after:{_TODAY_FILTER}"
            gn = _google_news(q_today, hl=hl, gl=gl, max_items=max_items)
            if gn:
                for r in gn:
                    r["category"] = category
                print(f"  [GN-TODAY] Got {len(gn)} articles for: {q}")
                return gn
    
    # 3) Google News without date filter (fallback when today has no coverage yet)
    for q, hl, gl in gn_queries:
        gn = _google_news(q, hl=hl, gl=gl, max_items=max_items)
        if gn:
            for r in gn:
                r["category"] = category
            return gn
    
    # 4) Try Bing News (multiple queries until one works)
    for q, hl in bing_queries:
        bn = _bing_news(q, hl=hl, max_items=max_items)
        if bn:
            for r in bn:
                r["category"] = category
            return bn
    
    return []


# ===== 各分类信息源采集（三级链路：主站→Google News→Bing News）=====

def fetch_brazil_economy():
    """巴西财经综合：Google News(葡/英) → Bing News(英)"""
    return _news_chain(
        primary_func=lambda: [],
        gn_queries=[
            ("Brasil economia mercado", "pt-BR", "BR"),
            ("Brazil economy business markets", "en", "US"),
        ],
        bing_queries=[
            ("Brazil economy business", "en"),
            ("Brasil economia", "pt-BR"),
        ],
        category="巴西财经综合",
        max_items=3,
    )

def _bcb_series(code):
    """获取BCB单一指标最新值（含前值用于对比变动方向）"""
    url = f"https://api.bcb.gov.br/dados/serie/bcdata.sgs.{code}/dados/ultimos/2?formato=json"
    text = _fetch(url, timeout=15)
    if text:
        try:
            data = json.loads(text)
            if data and len(data) > 0:
                latest = data[-1]   # 取最新值（列表末尾为最近一条）
                prev = data[-2] if len(data) >= 2 else None
                val = latest.get("valor", "").strip()
                date_s = latest.get("data", "")
                prev_val = prev.get("valor", "").strip() if prev else ""
                return val, date_s, prev_val
        except (json.JSONDecodeError, Exception):
            pass
    return "", "", ""

def fetch_bcb_focus():
    """巴西央行宏观经济指标（BCB公开API，多维度真实数据，含中文解读）"""
    results = []
    
    # === 扩展数据指标（全部来自BCB API，国内直连可用）===
    indicators = [
        # (指标名, BCB代码, 单位/说明)
        ("Selic基准利率", 432, "%"),
        ("IPCA通胀率(12个月)", 13522, "%"),
        ("GDP年增长率", 7326, "%"),
        ("美元/雷亚尔汇率(卖出)", 1, "BRL"),
        ("Ibovespa股指", 7, "点"),
        ("工业生产指数(同比)", 21860, "%"),
        ("零售销售指数(同比)", 1495, "%"),
        ("失业率(PNAD)", 24369, "%"),
        ("贸易顺差(月度)", 22707, "百万美元"),
        ("出口总额(月度)", 22708, "百万美元"),
        ("进口总额(月度)", 22709, "百万美元"),
        ("外国直接投资(月度)", 10838, "百万美元"),
        ("公共净债务/GDP", 4536, "%"),
    ]
    
    # 批量获取所有指标
    data = {}
    for name, code, unit in indicators:
        val, date_str, prev_val = _bcb_series(code)
        if val:
            data[name] = {"valor": val, "data": date_str, "unit": unit, "prev": prev_val}
    
    if not data:
        return results
    
    now_str = datetime.now(BJ_TZ).strftime("%Y-%m-%d")
    
    # === 类别1：货币政策与宏观 ===
    macro_items = []
    if "Selic基准利率" in data:
        s = data["Selic基准利率"]
        macro_items.append(f"Selic基准利率 {s['valor']}%（{s['data']}），是巴西核心货币政策利率，直接影响银行贷款、信用卡利率及整体信贷成本")
    if "IPCA通胀率(12个月)" in data:
        s = data["IPCA通胀率(12个月)"]
        macro_items.append(f"IPCA过去12个月累计通胀率 {s['valor']}%（参考 {s['data']}），IPCA是央行制定利率决策最核心的参照指标")
    if "GDP年增长率" in data:
        s = data["GDP年增长率"]
        macro_items.append(f"GDP年增长率 {s['valor']}%（{s['data']}），反映巴西经济活动整体扩张或收缩态势")
    if "失业率(PNAD)" in data:
        s = data["失业率(PNAD)"]
        macro_items.append(f"失业率 {s['valor']}%（{s['data']}），PNAD连续住户抽样调查数据，衡量劳动力市场松紧状况")
    
    if macro_items:
        title_parts = []
        if "Selic基准利率" in data:
            title_parts.append(f"Selic {data['Selic基准利率']['valor']}%")
        if "IPCA通胀率(12个月)" in data:
            title_parts.append(f"IPCA {data['IPCA通胀率(12个月)']['valor']}%")
        results.append({
            "title": "巴西货币政策关键指标 | " + " · ".join(title_parts),
            "summary": "。\n".join(macro_items) + "。",
            "source": "巴西中央银行 (BCB)",
            "category": "货币政策与宏观",
            "time": now_str,
            "url": "https://www.bcb.gov.br/estatisticas",
        })
    
    # === 类别2：贸易与国际收支 ===
    trade_items = []
    if "美元/雷亚尔汇率(卖出)" in data:
        s = data["美元/雷亚尔汇率(卖出)"]
        val_str = s['valor']
        change_note = ""
        if s.get('prev'):
            try:
                curr_f = float(val_str.replace(",", "."))
                prev_f = float(s['prev'].replace(",", "."))
                diff = curr_f - prev_f
                if abs(diff) > 0.0001:
                    # 汇率升高 = 雷亚尔贬值
                    arrow = "↑贬值" if diff > 0 else "↓升值"
                    change_note = f"，较上次{arrow} {abs(diff):.4f}"
            except Exception:
                pass
        trade_items.append(f"美元兑雷亚尔汇率 {val_str}{change_note}（{s['data']}），是影响巴西进出口价格竞争力及跨境投资回报率的核心变量")
    if "贸易顺差(月度)" in data:
        s = data["贸易顺差(月度)"]
        trade_items.append(f"月度贸易顺差 {s['valor']}百万美元（{s['data']}），体现巴西出口创汇能力与国际收支健康状况")
    if "出口总额(月度)" in data:
        s = data["出口总额(月度)"]
        trade_items.append(f"月度出口总额 {s['valor']}百万美元（{s['data']}），巴西主要出口品包括大豆、铁矿石、原油、肉类等大宗商品")
    if "进口总额(月度)" in data:
        s = data["进口总额(月度)"]
        trade_items.append(f"月度进口总额 {s['valor']}百万美元（{s['data']}），进口结构以机械设备、电子元器件、化工产品为主")
    if "外国直接投资(月度)" in data:
        s = data["外国直接投资(月度)"]
        trade_items.append(f"月度外国直接投资(FDI) {s['valor']}百万美元（{s['data']}），反映国际资本对巴西市场的信心与投入规模")
    
    if trade_items:
        title_val = ""
        if "贸易顺差(月度)" in data:
            title_val = f"顺差 {data['贸易顺差(月度)']['valor']}百万美元"
        results.append({
            "title": f"巴西对外贸易与投资数据 | {title_val}",
            "summary": "。\n".join(trade_items) + "。\n中国是巴西第一大贸易伙伴，上述进出口数据直接影响中巴双边贸易态势，建议关注大宗商品价格波动对巴西出口收入的传导效应。",
            "source": "巴西中央银行 (BCB)",
            "category": "贸易与国际收支",
            "time": now_str,
            "url": "https://www.bcb.gov.br/estatisticas/balancopagamentos",
        })
    
    # === 类别3：产业与市场 ===
    market_items = []
    if "Ibovespa股指" in data:
        s = data["Ibovespa股指"]
        val_str = s['valor']
        change_note = ""
        if s.get('prev'):
            try:
                curr_f = float(val_str.replace(",", "."))
                prev_f = float(s['prev'].replace(",", "."))
                diff = curr_f - prev_f
                if abs(diff) > 0.1:
                    pct = diff / prev_f * 100 if prev_f != 0 else 0
                    arrow = "↑" if diff > 0 else "↓"
                    change_note = f"，较前次{arrow}{abs(diff):.0f}点（{pct:+.2f}%）"
            except Exception:
                pass
        market_items.append(f"Ibovespa圣保罗股指报 {val_str}点{change_note}（{s['data']}），是拉丁美洲最重要的股票市场基准指数，涵盖巴西交易所最具流动性的蓝筹股")
    if "工业生产指数(同比)" in data:
        s = data["工业生产指数(同比)"]
        market_items.append(f"工业生产同比变动 {s['valor']}%（{s['data']}），衡量制造业、采矿业和公用事业产出变化，是判断经济周期阶段的重要先行指标")
    if "零售销售指数(同比)" in data:
        s = data["零售销售指数(同比)"]
        market_items.append(f"零售销售同比变动 {s['valor']}%（{s['data']}），反映居民消费意愿与购买力，是国内需求端的关键晴雨表")
    if "公共净债务/GDP" in data:
        s = data["公共净债务/GDP"]
        market_items.append(f"公共部门净债务占GDP比重 {s['valor']}%（{s['data']}），衡量政府财政可持续性与主权信用风险，是国际评级机构关注的核心财政指标")
    
    if market_items:
        ibov_val = data.get("Ibovespa股指", {}).get("valor", "")
        results.append({
            "title": f"巴西产业与金融市场数据 | Ibovespa {ibov_val}点",
            "summary": "。\n".join(market_items) + "。",
            "source": "巴西中央银行 / IBGE (BCB)",
            "category": "产业与市场",
            "time": now_str,
            "url": "https://www.bcb.gov.br/estatisticas",
        })
    
    return results

def _try_ibge_rss():
    """尝试从IBGE RSS获取新闻"""
    results = []
    rss_urls = [
        "https://agenciadenoticias.ibge.gov.br/en/component/ninjarsssyndicator/?feed_id=1&format=raw",
        "https://agenciadenoticias.ibge.gov.br/agencia-noticias/ultimas.html?format=feed&type=rss",
    ]
    text = ""
    for url in rss_urls:
        text = _fetch(url, headers={"Accept": "application/rss+xml, application/xml, text/xml"}, timeout=15)
        if text and "<item>" in text:
            break
    
    if text and "<item>" in text:
        items = text.split("<item>")[1:]
        for item in items[:3]:
            title = _extract(item, "<title>", "</title>")
            link = _extract(item, "<link>", "</link>")
            pubdate = _extract(item, "<pubDate>", "</pubDate>")
            desc = _extract(item, "<description>", "</description>")
            if title and link:
                results.append({
                    "title": _clean(title),
                    "summary": _clean(desc)[:400] if desc else _clean(title),
                    "source": "IBGE 巴西统计局",
                    "category": "宏观经济数据",
                    "time": (pubdate or "")[:25],
                    "url": link,
                })
    return results

def fetch_ibge():
    """IBGE 巴西统计局：RSS直连 → Bing News降级 + IPCA API（多重保障）"""
    results = _try_ibge_rss()
    
    # 如果 RSS 失败，用 Bing News 搜索巴西经济数据相关新闻
    if not results:
        bn = _bing_news("Brazil economic data IBGE inflation", hl="en", max_items=3)
        for r in bn:
            r["category"] = "宏观经济数据"
        results = bn
    
    # IPCA API（独立获取，总有数据）
    try:
        ipca_text = _fetch("https://servicodados.ibge.gov.br/api/v3/agregados/7060/periodos/-1/variaveis?localidades=N1[all]", timeout=15)
        if ipca_text:
            ipca_data = json.loads(ipca_text)
            if ipca_data and isinstance(ipca_data, list) and len(ipca_data) > 0:
                result_item = ipca_data[0]
                resultados = result_item.get("resultados", [])
                if resultados:
                    s = resultados[0].get("series", [])
                    if s:
                        vals = s[0].get("serie", {})
                        sorted_dates = sorted(vals.keys(), reverse=True)
                        if sorted_dates:
                            latest_val = vals[sorted_dates[0]]
                            results.append({
                                "title": "巴西IPCA月度通胀率更新",
                                "summary": f"最新IPCA月度通胀率为 {latest_val}%（参考月份 {sorted_dates[0]}），IPCA是巴西央行衡量通胀的核心指标，直接影响Selic利率决策方向。",
                                "source": "IBGE 巴西统计局",
                                "category": "宏观经济数据",
                                "time": sorted_dates[0],
                                "url": "https://www.ibge.gov.br/explica/inflacao.php",
                            })
    except Exception:
        pass
    
    return results

def _try_xinhua_brazil():
    """尝试从新华网葡语版获取巴西相关新闻"""
    results = []
    xinhua_urls = [
        "https://portuguese.news.cn/america-latina/",
        "https://portuguese.news.cn/",
    ]
    for url in xinhua_urls:
        text = _fetch(url, timeout=15)
        if text and len(text) > 500:
            links = re.findall(r'href="(https?://portuguese\.news\.cn/[^"]{10,})"', text)
            titles_raw = re.findall(r'(?:title|alt)="([^"]{10,120})"', text)
            seen_titles = set()
            for t, l in zip(titles_raw[:5], links[:5]):
                clean_t = _clean(t)
                if clean_t and clean_t not in seen_titles and len(clean_t) > 15:
                    seen_titles.add(clean_t)
                    results.append({
                        "title": clean_t,
                        "summary": clean_t[:400],
                        "source": "新华网葡语版",
                        "category": "中巴关系与政经",
                        "time": datetime.now(BJ_TZ).strftime("%Y-%m-%d"),
                        "url": l,
                    })
            if results:
                break
    return results

def fetch_china_brazil():
    """中巴关系与政经：新华网 → Google News(中/葡) → Bing News(中/英)"""
    return _news_chain(
        primary_func=_try_xinhua_brazil,
        gn_queries=[
            ("中国 巴西 经贸 合作", "zh-CN", "CN"),
            ("China Brasil relações comércio", "pt-BR", "BR"),
        ],
        bing_queries=[
            ("中国 巴西 经贸", "zh-CN"),
            ("China Brazil trade relations", "en"),
        ],
        category="中巴关系与政经",
        max_items=3,
    )

def fetch_international_brazil():
    """国际视角巴西：Google News(英) → Bing News(英)"""
    return _news_chain(
        primary_func=lambda: [],
        gn_queries=[
            ("Brazil economy markets international", "en", "US"),
        ],
        bing_queries=[
            ("Brazil economy markets", "en"),
            ("Brasil economia internacional", "pt-BR"),
        ],
        category="国际视角巴西",
        max_items=3,
    )

def _try_nanmei():
    """尝试从南美侨报网获取新闻"""
    results = []
    try:
        text = _fetch("https://www.br-cn.com/", timeout=15)
        if text and len(text) > 500:
            titles_raw = re.findall(r'(?:title|alt)="([^"]{8,80})"', text)
            links = re.findall(r'href="(https?://www\.br-cn\.com[^"]{5,})"', text)
            seen = set()
            for t, l in zip(titles_raw, links[:5]):
                clean_t = _clean(t)
                if clean_t and clean_t not in seen and len(clean_t) > 12:
                    seen.add(clean_t)
                    results.append({
                        "title": clean_t,
                        "summary": clean_t[:400],
                        "source": "南美侨报网",
                        "category": "跨境贸易与华人",
                        "time": datetime.now(BJ_TZ).strftime("%Y-%m-%d"),
                        "url": l,
                    })
    except Exception:
        pass
    return results

def fetch_chinese_community():
    """跨境贸易与华人：南美侨报网 → Google News(中/葡) → Bing News(中/英)"""
    return _news_chain(
        primary_func=_try_nanmei,
        gn_queries=[
            ("巴西 华人 贸易 清关", "zh-CN", "CN"),
            ("comunidade chinesa Brasil comércio", "pt-BR", "BR"),
        ],
        bing_queries=[
            ("巴西 华人 贸易 关税", "zh-CN"),
            ("Chinese community Brazil trade", "en"),
        ],
        category="跨境贸易与华人",
        max_items=3,
    )

def _try_receita():
    """尝试从巴西税务局官网抓取当日真实新闻并翻译为中文。
    翻译成功才返回；否则返回空列表，让链路交给 Google News 提供每日新鲜内容。"""
    today_str = datetime.now(BJ_TZ).strftime("%Y-%m-%d")

    try:
        text = _fetch("https://www.gov.br/receitafederal/pt-br/assuntos/noticias", timeout=15)
        if text and len(text) > 500:
            items_found = re.findall(r'(?:title|alt)="([^"]{15,120})"', text)
            links_found = re.findall(r'href="(/receitafederal[^"]{10,})"', text)
            seen_t = set()
            raw_arts = []
            skip_kw = ["receita federal", "gov.br", "javascript", "toggle", "menu",
                       "search", "buscar", "acessibilidade", "governo"]
            for t, l in zip(items_found[:10], links_found[:10]):
                ct = _clean(t)
                if ct and ct not in seen_t and len(ct) > 18:
                    if not any(kw in ct.lower() for kw in skip_kw):
                        seen_t.add(ct)
                        raw_arts.append((ct, "https://www.gov.br" + l if l.startswith("/") else l))
            if raw_arts:
                titles_pt = "；".join([a[0] for a in raw_arts[:4]])
                zh_titles = translate_to_chinese(titles_pt)
                if is_chinese(zh_titles):
                    return [{
                        "title": f"巴西联邦税务局今日公告（{today_str}）",
                        "summary": zh_titles,
                        "source": "巴西联邦税务局 (Receita Federal)",
                        "category": "税务与合规",
                        "time": today_str,
                        "url": raw_arts[0][1],
                    }]
    except Exception:
        pass

    # 官网不可达或翻译失败 → 返回空列表，交给 Google News 提供当天新鲜内容
    return []

def fetch_tax_compliance():
    """税务与合规：巴西税务局 → Google News(葡/英) → Bing News(英)"""
    return _news_chain(
        primary_func=_try_receita,
        gn_queries=[
            ("Receita Federal Brasil imposto", "pt-BR", "BR"),
            ("Brazil tax regulation customs", "en", "US"),
        ],
        bing_queries=[
            ("Brazil tax customs regulation", "en"),
            ("Receita Federal Brasil", "pt-BR"),
        ],
        category="税务与合规",
        max_items=3,
    )

# ===== 主流程 =====
def main():
    print("=" * 50)
    print("[START] Brazil News Daily Bot")
    
    if not AUTH_CODE:
        print("[FAIL] QQ_MAIL_AUTH_CODE not set. Abort.")
        sys.exit(1)
    
    bj_now = datetime.now(BJ_TZ)
    bj_today = bj_now.date()
    print(f"[TIME] {bj_now.strftime('%Y-%m-%d %H:%M:%S')} BJT  |  Filter date: {bj_today}")
    
    # 设置全局今日日期过滤（供 _news_chain 内 Google News after: 使用）
    global _TODAY_FILTER
    _TODAY_FILTER = bj_today.strftime("%Y-%m-%d")
    print(f"[FILTER] Google News date constraint: after:{_TODAY_FILTER}")
    
    # Collect from all sources
    all_articles = []
    
    print("[FETCH] 巴西财经综合 (Google News)...")
    all_articles.extend(fetch_brazil_economy())
    
    print("[FETCH] 巴西央行宏观经济指标 (BCB API)...")
    all_articles.extend(fetch_bcb_focus())
    
    print("[FETCH] IBGE 统计局数据...")
    all_articles.extend(fetch_ibge())
    
    print("[FETCH] 中巴关系与政经 (新华网+Google News)...")
    all_articles.extend(fetch_china_brazil())
    
    print("[FETCH] 国际视角巴西 (Google News国际源)...")
    all_articles.extend(fetch_international_brazil())
    
    print("[FETCH] 跨境贸易与华人 (南美侨报+Google News)...")
    all_articles.extend(fetch_chinese_community())
    
    print("[FETCH] 税务与合规 (Receita Federal+Google News)...")
    all_articles.extend(fetch_tax_compliance())
    
    # === FILTER: only today's articles ===
    print(f"[FILTER] Raw articles: {len(all_articles)} → keeping today ({bj_today}) only...")
    today_articles = [a for a in all_articles if is_today_article(a, bj_today)]
    excluded = len(all_articles) - len(today_articles)
    if excluded > 0:
        print(f"[FILTER] Excluded {excluded} non-today/fallback articles")
    print(f"[FILTER] Today articles: {len(today_articles)}")
    
    # Deduplicate by title similarity
    seen = set()
    deduped = []
    for a in today_articles:
        key = a["title"][:30].lower().strip()
        if key not in seen:
            seen.add(key)
            deduped.append(a)
    
    # === 中文翻译 ===
    print(f"[TRANSLATE] Translating non-Chinese articles to Chinese...")
    for a in deduped:
        translate_article(a)
    
    src_count = len(set(a["source"] for a in deduped))
    print(f"[STATS] {len(deduped)} unique today articles from {src_count} sources")
    
    if len(deduped) == 0:
        print("[WARN] No today articles found. Skipping email and WeChat publish.")
        print("[DONE] === Complete (no content today) ===")
        return
    
    # Build HTML body
    html = build_html(deduped, bj_now)
    
    # Build GitHub Pages page (国内可访问)
    print("[PAGES] Building GitHub Pages HTML...")
    pages_html = build_github_page(deduped, bj_now)
    pages_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_pages")
    os.makedirs(pages_dir, exist_ok=True)
    pages_path = os.path.join(pages_dir, "index.html")
    with open(pages_path, "w", encoding="utf-8") as f:
        f.write(pages_html)
    print(f"[PAGES] Saved -> {pages_path}")
    
    # Build Word document
    print("[DOCX] Generating Word attachment...")
    doc = build_docx(deduped, bj_now)
    tmp_docx = os.path.join(tempfile.gettempdir(), f"brazil_news_{bj_now.strftime('%Y%m%d')}.docx")
    doc.save(tmp_docx)
    print(f"[DOCX] Saved -> {tmp_docx}")
    
    # Build WeChat publish HTML
    print("[HTML] Generating WeChat publish HTML...")
    wechat_html = build_wechat_publish_html(deduped, bj_now)
    tmp_html = os.path.join(tempfile.gettempdir(), f"wechat_publish_{bj_now.strftime('%Y%m%d')}.html")
    with open(tmp_html, "w", encoding="utf-8") as f:
        f.write(wechat_html)
    print(f"[HTML] Saved -> {tmp_html}")
    
    # Locate cover image
    script_dir = os.path.dirname(os.path.abspath(__file__))
    cover_path = os.path.join(script_dir, "cover.png")
    cover_exists = os.path.exists(cover_path)
    if cover_exists:
        print(f"[COVER] Found cover image: {cover_path}")
    else:
        print("[COVER] WARN: cover.png not found, will skip cover attachment")
    
    # Build email with HTML body + 3 attachments (DOCX + Cover PNG + WeChat HTML)
    docx_filename = f"巴西资讯日报_{bj_now.strftime('%Y%m%d')}.docx"
    
    msg = EmailMessage()
    msg["From"] = SENDER
    msg["To"] = RECEIVER
    msg["Subject"] = f"巴西资讯日报 | {bj_now.strftime('%Y-%m-%d')}"
    
    # Plain text + HTML alternative content
    msg.set_content("巴西资讯日报 - 请使用支持HTML的邮件客户端查看。\n附件：Word文档、公众号HTML文件、封面图片。")
    msg.add_alternative(html, subtype="html")
    
    # Attach Word document
    with open(tmp_docx, "rb") as f:
        docx_data = f.read()
    msg.add_attachment(
        docx_data,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=docx_filename,
    )
    
    # Attach WeChat publish HTML
    html_filename = f"公众号粘贴_{bj_now.strftime('%Y%m%d')}.html"
    with open(tmp_html, "rb") as f:
        html_data = f.read()
    msg.add_attachment(
        html_data,
        maintype="text",
        subtype="html",
        filename=html_filename,
    )
    
    # Attach cover image
    if cover_exists:
        with open(cover_path, "rb") as f:
            cover_data = f.read()
        msg.add_attachment(
            cover_data,
            maintype="image",
            subtype="png",
            filename="巴西资讯日报_封面.png",
        )
        print("[COVER] Cover image attached to email")
    
    print("[SEND] Connecting to QQ SMTP...")
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, context=ctx) as server:
            server.login(SENDER, AUTH_CODE)
            server.send_message(msg)
        print(f"[OK] Email sent successfully! {len(deduped)} articles + DOCX attachment -> {RECEIVER}")
    except Exception as e:
        print(f"[FAIL] Email send error: {e}")
        sys.exit(1)
    
    # === WeChat auto-publish (optional) ===
    wechat_publish(deduped, bj_now)
    
    # Clean up temp files
    for tmp_file in [tmp_docx, tmp_html]:
        try:
            os.remove(tmp_file)
        except OSError:
            pass
    
    print("[DONE] === Complete ===")

if __name__ == "__main__":
    main()
